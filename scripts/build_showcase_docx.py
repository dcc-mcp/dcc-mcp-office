"""Build the polished Word source artifact used by the showcase gallery."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CONTENT_WIDTH_DXA = 10080
ACCENT = "3157F6"
CYAN = "70D7FF"
CORAL = "FF6B5E"
NAVY = "0B1020"
INK = "172033"
MUTED = "5E6678"
PALE = "F4F2ED"
PALE_BLUE = "E9F2FF"
WHITE = "FFFFFF"


def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, *, top: int = 100, bottom: int = 100, start: int = 120, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell: Any, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: list[int], *, indent_dxa: int = 0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        node = properties.find(qn(tag))
        if node is not None:
            properties.remove(node)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    properties.append(table_width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa, strict=True):
            set_cell_width(cell, width_dxa)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table: Any, *, color: str = "D0D5DD", size: int = 4, inside: bool = True) -> None:
    properties = table._tbl.tblPr
    borders = properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names.extend(["insideH", "insideV"])
    for name in names:
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_paragraph_rule(paragraph: Any, color: str = ACCENT, size: int = 12) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_run_font(
    run: Any,
    name: str = "Aptos",
    size: float | None = None,
    color: str | None = None,
    *,
    display: bool = False,
) -> None:
    if display:
        name = "Aptos Display"
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 13, ACCENT, 11, 5),
        ("Heading 3", 11.5, ACCENT, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section: Any) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.66)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)


def add_header_footer(document: DocumentType) -> None:
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "DCC-MCP  /  OFFICE AUTOMATION  /  EVIDENCE BRIEF"
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, size=7.8, color=MUTED)
        run.font.bold = True
        run.font.all_caps = True

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.paragraph_format.space_before = Pt(0)
    left = paragraph.add_run("EDITABLE SOURCE  ·  NATIVE OFFICE PROOF")
    set_run_font(left, size=7.8, color=MUTED)
    separator = paragraph.add_run("\t")
    set_run_font(separator, size=7.8, color=MUTED)
    add_page_field(paragraph)


def add_masthead(document: DocumentType, content: dict[str, Any]) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run(content["kicker"])
    set_run_font(run, size=8, color=ACCENT)
    run.font.bold = True
    run.font.all_caps = True

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(content["title"])
    set_run_font(run, size=31, color=NAVY, display=True)
    run.font.bold = True

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run(content["subtitle"])
    set_run_font(run, size=11.2, color=MUTED)

    metadata = content["metadata"]
    row = document.add_paragraph()
    row.paragraph_format.space_after = Pt(5)
    for index, item in enumerate(
        (
            f"Prepared for  {metadata['prepared_for']}",
            f"Status  {metadata['status']}",
            f"Evidence  {metadata['evidence_date']}",
        )
    ):
        if index:
            separator = row.add_run("   |   ")
            set_run_font(separator, size=7.7, color="98A2B3")
        run = row.add_run(item)
        set_run_font(run, size=7.7, color=MUTED)
        run.font.bold = index == 1
    set_paragraph_rule(row)


def add_banner_image(document: DocumentType, image_path: Path) -> None:
    if not image_path.is_file():
        raise FileNotFoundError(f"showcase banner not found: {image_path}")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(7.0), height=Inches(1.72))
    inline = run._r.xpath(".//wp:inline")
    if inline:
        doc_properties = inline[0].xpath("./wp:docPr")
        if doc_properties:
            doc_properties[0].set("title", "Governed Office evidence")
            doc_properties[0].set("descr", "Editorial still life of editable Office outputs and verification evidence")


def add_decision_callout(document: DocumentType, decision: dict[str, str]) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [7920, 2160])
    cell, metric_cell = table.rows[0].cells
    set_cell_shading(cell, NAVY)
    set_cell_shading(metric_cell, ACCENT)
    set_cell_margins(cell, top=140, bottom=140, start=190, end=190)
    set_cell_margins(metric_cell, top=140, bottom=140, start=120, end=120)
    label = cell.paragraphs[0]
    label.paragraph_format.space_after = Pt(5)
    run = label.add_run(decision["label"])
    set_run_font(run, size=7.7, color=CYAN)
    run.font.bold = True
    title = cell.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run(decision["title"])
    set_run_font(run, size=13.5, color=WHITE, display=True)
    run.font.bold = True
    body = cell.add_paragraph()
    body.paragraph_format.space_after = Pt(0)
    run = body.add_run(decision["body"])
    set_run_font(run, size=9.1, color="E4E7EC")
    metric = metric_cell.paragraphs[0]
    metric.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metric.paragraph_format.space_after = Pt(0)
    value = metric.add_run("11")
    set_run_font(value, size=29, color=WHITE, display=True)
    value.font.bold = True
    metric.add_run("\n")
    caption = metric.add_run("LIVE\nCAPABILITIES")
    set_run_font(caption, size=7.4, color=WHITE)
    caption.font.bold = True
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(2)


def add_capability_table(document: DocumentType, capabilities: list[dict[str, str]]) -> None:
    heading = document.add_paragraph("Runtime capability ledger", style="Heading 1")
    heading.paragraph_format.space_before = Pt(5)
    intro = document.add_paragraph(
        "Live handshakes advertise only the application-specific operations below; counts are evidence, not projections."
    )
    intro.paragraph_format.space_after = Pt(5)
    for run in intro.runs:
        set_run_font(run, size=8.8, color=MUTED)

    table = document.add_table(rows=1, cols=4)
    widths = [1840, 760, 2460, 5020]
    set_table_geometry(table, widths)
    set_table_borders(table)
    for cell, label in zip(table.rows[0].cells, ("Application", "Live", "Provider", "Verified work"), strict=True):
        set_cell_shading(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_run_font(run, size=8.2, color=WHITE)
        run.font.bold = True

    for item in capabilities:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=72, bottom=72, start=110, end=110)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        values = (item["application"], item["count"], item["mode"], item["proof"])
        for index, (cell, value) in enumerate(zip(cells, values, strict=True)):
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if index == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_run_font(run, size=8.8, color=ACCENT if index == 1 else INK, display=index == 1)
            run.font.bold = index in (0, 1)


def add_workflow(document: DocumentType, workflow: list[dict[str, str]]) -> None:
    section_number = document.add_paragraph()
    section_number.paragraph_format.space_after = Pt(2)
    run = section_number.add_run("02  /  EXECUTION SYSTEM")
    set_run_font(run, size=7.8, color=CORAL)
    run.font.bold = True
    document.add_paragraph("Governed execution flow", style="Heading 1")
    intro = document.add_paragraph(
        "The suite treats automation as a contract: discover first, make risk visible, mutate only with authority, then prove the result."
    )
    intro.paragraph_format.space_after = Pt(7)
    for run in intro.runs:
        set_run_font(run, size=9.6, color=MUTED)

    table = document.add_table(rows=0, cols=2)
    widths = [1160, 8920]
    set_table_geometry(table, widths)
    set_table_borders(table, color="D6E4F0", inside=True)
    for index, item in enumerate(workflow):
        cells = table.add_row().cells
        for cell, width in zip(cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=92, bottom=92, start=120, end=120)
        set_cell_shading(cells[0], ACCENT if index < 4 else CORAL)
        number = cells[0].paragraphs[0]
        number.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number.paragraph_format.space_after = Pt(0)
        run = number.add_run(item["step"])
        set_run_font(run, size=11, color=WHITE, display=True)
        run.font.bold = True
        if index % 2 == 0:
            set_cell_shading(cells[1], "F8FAFC")
        paragraph = cells[1].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        title = paragraph.add_run(item["name"] + "  ")
        set_run_font(title, size=9.5, color=NAVY, display=True)
        title.font.bold = True
        body = paragraph.add_run(item["body"])
        set_run_font(body, size=8.8, color=INK)


def add_safeguards(document: DocumentType, safeguards: list[dict[str, str]]) -> None:
    document.add_paragraph("Operational safeguards", style="Heading 1")
    table = document.add_table(rows=1, cols=3)
    widths = [3360, 3360, 3360]
    set_table_geometry(table, widths)
    for index, (cell, item) in enumerate(zip(table.rows[0].cells, safeguards, strict=True)):
        set_cell_width(cell, widths[index])
        set_cell_margins(cell, top=125, bottom=125, start=135, end=135)
        set_cell_shading(cell, PALE_BLUE if index == 1 else PALE)
        title = cell.paragraphs[0]
        title.paragraph_format.space_after = Pt(5)
        run = title.add_run(item["title"])
        set_run_font(run, size=9.5, color=ACCENT, display=True)
        run.font.bold = True
        body = cell.add_paragraph()
        body.paragraph_format.space_after = Pt(0)
        run = body.add_run(item["body"])
        set_run_font(run, size=8.5, color=INK)


def add_boundary(document: DocumentType, boundary: str) -> None:
    heading = document.add_paragraph("Evidence boundary", style="Heading 2")
    heading.paragraph_format.space_before = Pt(9)
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="D0D5DD", inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F9FAFB")
    set_cell_margins(cell, top=105, bottom=105, start=145, end=145)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(boundary)
    set_run_font(run, size=8.6, color=MUTED)


def build_document(content: dict[str, Any], output: Path) -> None:
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    add_header_footer(document)
    properties = document.core_properties
    properties.title = content["title"]
    properties.subject = "dcc-mcp-office verified showcase"
    properties.author = "DCC-MCP"
    properties.keywords = "Office automation, PowerPoint, Word, Excel, evidence"

    add_masthead(document, content)
    add_banner_image(document, output.parent / "assets" / "document-evidence-banner.jpg")
    add_decision_callout(document, content["decision"])
    add_capability_table(document, content["capabilities"])

    break_paragraph = document.add_paragraph()
    break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    add_workflow(document, content["workflow"])
    add_safeguards(document, content["safeguards"])
    add_boundary(document, content["boundary"])

    # Keep both pages in the same section so the masthead's quiet header/footer repeats.
    if len(document.sections) != 1:
        raise RuntimeError("showcase document unexpectedly contains multiple sections")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def parse_args() -> argparse.Namespace:
    root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--input",
        type=Path,
        default=root / "showcase" / "word-executive-brief" / "content.json",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=root / "showcase" / "word-executive-brief" / "dcc-mcp-office-executive-brief.docx",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    content = json.loads(args.input.read_text(encoding="utf-8"))
    build_document(content, args.output)
    print(json.dumps({"artifact": str(args.output.resolve()), "pages_expected": 2}, ensure_ascii=False))


if __name__ == "__main__":
    main()
